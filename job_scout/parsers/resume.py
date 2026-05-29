"""Resume text extraction for PDF / DOCX / TXT.

Dispatches by extension. Optional dependencies (`pypdf`, `python-docx`) are
imported lazily so importing this module never fails just because a parser
backend isn't installed.
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path


class UnsupportedResumeFormat(ValueError):
    """Raised when the uploaded file's extension is not supported."""


def extract_resume_text(filename: str, data: bytes) -> str:
    """Best-effort plain-text extraction from a resume upload.

    Dispatches on file extension. Raises `UnsupportedResumeFormat` for
    anything other than .pdf, .docx, .txt.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return _extract_txt(data)
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    raise UnsupportedResumeFormat(
        f"Unsupported resume format '{suffix}'. Supported: .pdf, .docx, .txt"
    )


def _extract_txt(data: bytes) -> str:
    # Tolerate UTF-8 with BOM and common single-byte fallbacks.
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf is required for PDF resumes. Install with `pip install pypdf`."
        ) from e

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX resumes. "
            "Install with `pip install python-docx`."
        ) from e

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Tables: many executive resumes use tables for layout.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)
